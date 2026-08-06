import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from io import BytesIO

class ReportExporter:
    def export_pdf(self, case_title: str, case_number: str, summary_text: str, evidence_list: list, suspect_list: list, witness_list: list, timeline_events: list) -> BytesIO:
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()
        
        # Modify existing styles safely or add custom ones
        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=24,
            leading=28,
            textColor=colors.HexColor('#0f172a'), # Slate-900
            spaceAfter=20
        )
        
        h2_style = ParagraphStyle(
            'ReportH2',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            textColor=colors.HexColor('#1e3a8a'), # Blue-900
            spaceBefore=15,
            spaceAfter=10,
            keepWithNext=True
        )

        body_style = ParagraphStyle(
            'ReportBody',
            parent=styles['BodyText'],
            fontName='Helvetica',
            fontSize=10,
            leading=15,
            textColor=colors.HexColor('#334155'), # Slate-700
            spaceAfter=8
        )

        bullet_style = ParagraphStyle(
            'ReportBullet',
            parent=styles['Bullet'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#334155'),
            spaceAfter=5
        )

        story = []

        # Title
        story.append(Paragraph("VISHAM CRIMINAL INVESTIGATION REPORT", title_style))
        story.append(Spacer(1, 10))

        # Case Meta Box Info
        meta_data = [
            [Paragraph("<b>Case Number:</b>", body_style), Paragraph(case_number, body_style),
             Paragraph("<b>Case Title:</b>", body_style), Paragraph(case_title, body_style)]
        ]
        meta_table = Table(meta_data, colWidths=[100, 150, 100, 150])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f1f5f9')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#cbd5e1')),
            ('TOPPADDING', (0,0), (-1,-1), 10),
            ('BOTTOMPADDING', (0,0), (-1,-1), 10),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 20))

        # 1. Executive Summary
        story.append(Paragraph("1. Executive Summary", h2_style))
        story.append(Paragraph(summary_text.replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 15))

        # 2. Evidence Registry
        story.append(Paragraph("2. Evidence Analysis", h2_style))
        if not evidence_list:
            story.append(Paragraph("No physical or digital evidence registered yet.", body_style))
        else:
            evidence_data = [["ID", "Title", "Type", "Uploaded By"]]
            for ev in evidence_list:
                evidence_data.append([
                    str(ev.id),
                    ev.title,
                    ev.evidence_type,
                    ev.uploaded_by
                ])
            ev_table = Table(evidence_data, colWidths=[40, 180, 140, 140])
            ev_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e40af')),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f8fafc')]),
                ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#e2e8f0')),
                ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ]))
            story.append(ev_table)
        story.append(Spacer(1, 15))

        # 3. Suspect Analysis
        story.append(Paragraph("3. Suspect Analysis", h2_style))
        if not suspect_list:
            story.append(Paragraph("No suspect profiles associated with this case yet.", body_style))
        else:
            suspect_data = [["Name", "Alias", "Risk Level", "Status"]]
            for s in suspect_list:
                suspect_data.append([
                    s.full_name,
                    s.alias or "N/A",
                    s.risk_level,
                    s.status
                ])
            sus_table = Table(suspect_data, colWidths=[150, 120, 110, 120])
            sus_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#991b1b')),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f8fafc')]),
                ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#e2e8f0')),
                ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ]))
            story.append(sus_table)
        story.append(Spacer(1, 15))

        # 4. Witness Profiles
        story.append(Paragraph("4. Witness Analysis", h2_style))
        if not witness_list:
            story.append(Paragraph("No witnesses interview records registered yet.", body_style))
        else:
            for w in witness_list:
                story.append(Paragraph(f"<b>Witness Name:</b> {w.full_name} | <b>Credibility:</b> {w.credibility}", body_style))
                story.append(Paragraph(f"<b>Statement:</b> {w.statement}", ParagraphStyle('Stmt', parent=body_style, leftIndent=20, fontName='Helvetica-Oblique')))
                story.append(Spacer(1, 5))
        story.append(Spacer(1, 15))

        # 5. Timeline Reconstruction
        story.append(Paragraph("5. Timeline Reconstruction", h2_style))
        if not timeline_events:
            story.append(Paragraph("No timeline events plotted yet.", body_style))
        else:
            for t in timeline_events:
                time_str = t.event_time.strftime('%Y-%m-%d %H:%M') if hasattr(t.event_time, 'strftime') else str(t.event_time)
                story.append(Paragraph(f"• <b>{time_str}</b>: {t.title} - {t.description}", bullet_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer

    def export_docx(self, case_title: str, case_number: str, summary_text: str, evidence_list: list, suspect_list: list, witness_list: list, timeline_events: list) -> BytesIO:
        doc = Document()
        doc.add_heading("VISHAM CRIMINAL INVESTIGATION REPORT", level=0)

        # Meta
        doc.add_paragraph(f"Case Number: {case_number}")
        doc.add_paragraph(f"Case Title: {case_title}")
        doc.add_paragraph()

        # Summary
        doc.add_heading("1. Executive Summary", level=1)
        doc.add_paragraph(summary_text)

        # Evidence
        doc.add_heading("2. Evidence Analysis", level=1)
        if not evidence_list:
            doc.add_paragraph("No evidence registered.")
        else:
            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Title'
            hdr_cells[2].text = 'Type'
            hdr_cells[3].text = 'Uploaded By'
            for ev in evidence_list:
                row_cells = table.add_row().cells
                row_cells[0].text = str(ev.id)
                row_cells[1].text = ev.title
                row_cells[2].text = ev.evidence_type
                row_cells[3].text = ev.uploaded_by

        doc.add_paragraph()

        # Suspects
        doc.add_heading("3. Suspect Analysis", level=1)
        if not suspect_list:
            doc.add_paragraph("No suspects profiles associated.")
        else:
            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Alias'
            hdr_cells[2].text = 'Risk Level'
            hdr_cells[3].text = 'Status'
            for s in suspect_list:
                row_cells = table.add_row().cells
                row_cells[0].text = s.full_name
                row_cells[1].text = s.alias or "N/A"
                row_cells[2].text = s.risk_level
                row_cells[3].text = s.status

        doc.add_paragraph()

        # Witnesses
        doc.add_heading("4. Witness Analysis", level=1)
        if not witness_list:
            doc.add_paragraph("No witnesses associated.")
        else:
            for w in witness_list:
                doc.add_paragraph(f"Witness: {w.full_name} (Credibility: {w.credibility})")
                doc.add_paragraph(f"Statement: \"{w.statement}\"")
                doc.add_paragraph()

        # Timeline
        doc.add_heading("5. Timeline Reconstruction", level=1)
        if not timeline_events:
            doc.add_paragraph("No timeline events.")
        else:
            for t in timeline_events:
                time_str = t.event_time.strftime('%Y-%m-%d %H:%M') if hasattr(t.event_time, 'strftime') else str(t.event_time)
                doc.add_paragraph(f"{time_str} - {t.title}: {t.description}", style='List Bullet')

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
